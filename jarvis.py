#!/usr/bin/env python3
# ============================================================
# Jarvis Prime v6 — Full Cognitive Assistant (local-first)
# ============================================================
# Default STT quality: HIGH (Whisper "medium.en")
# Key features:
#  - Robust STT: energy VAD (+ optional WebRTC VAD), faster-whisper
#  - Single‑word "open" intents: "google", "gmail", "youtube", etc.
#  - Smart routing: factual → web; math → solver; chit‑chat → funny LLM
#  - British voice auto-pick on Windows (if available)
#  - Minimal console UX: prints your input + "Thinking…" spinner
# ------------------------------------------------------------

import os, sys, time, json, math, re, shutil, zipfile, queue, threading, tempfile, traceback, webbrowser, logging, logging.handlers, subprocess, wave, difflib
from datetime import datetime, timedelta
from typing import Optional, Dict, Any, List, Tuple, Deque
from collections import deque

# Force CPU Whisper unless user sets GPU; keep logs quiet
os.environ.setdefault("CUDA_VISIBLE_DEVICES", "-1")
os.environ.setdefault("TF_CPP_MIN_LOG_LEVEL", "3")
os.environ.setdefault("PYTORCH_ENABLE_MPS_FALLBACK", "1")

# Third‑party deps
import numpy as np
import requests
import sounddevice as sd
import psutil
import pyperclip

# Optional VAD
try:
    import webrtcvad
    WEBRTC_OK = True
except Exception:
    WEBRTC_OK = False

from faster_whisper import WhisperModel

# Office/PDF
from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
import pdfplumber
from PyPDF2 import PdfReader, PdfWriter

# Math
import sympy as sp

# Optional DOCX→PDF helper
DOCX2PDF_AVAILABLE = False
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

# ------------------------------------------------------------
# Configuration
# ------------------------------------------------------------
CONFIG: Dict[str, Any] = {
    "sample_rate": 16000,

    # ===== STT default: HIGH =====
    # You can switch at runtime: stt quality fast|balanced|high|best
    "whisper_model": "medium.en",
    "whisper_beam_size": 5,
    "whisper_temperature": 0.0,
    "whisper_language": "en",

    # VAD (tuned for snappy starts/stops)
    "webrtc_vad": True,               # requires webrtcvad (optional)
    "webrtc_vad_aggressiveness": 2,   # 0-3
    "vad_silence_ms": 300,
    "vad_pre_speech_ms": 450,
    "vad_max_utterance_s": 12,
    "vad_min_utterance_s": 0.30,

    # No wake word (always‑on push‑to‑talk style)
    "wake_word_enabled": False,
    "wake_words": ["hey jarvis", "jarvis"],

    # Keys / endpoints
    "serpapi_key": os.getenv("SERPAPI_KEY", ""),
    "ollama_api": "http://127.0.0.1:11434/api/generate",
    "ollama_model": "gemma3:4b",
    "ollama_timeout_s": 120,

    # Audio
    "input_device": None,             # set via `audio in <index>`

    # Files / state
    "memory_file": "jarvis_memory.json",
    "notes_file": "jarvis_notes.json",
    "config_file": "jarvis_config.json",
    "log_file": "jarvis.log",

    # TTS
    "voice_rate": 2,                  # SAPI -10..+10 (2 = a bit faster)
    "voice_volume": 100,              # SAPI 0..100
    "voice_index": 0,
    "reply_language": "en",           # LLM reply language

    # Personas
    "system_style_concise": "You are Jarvis. Reply in ≤2 short sentences. Be crisp, direct, useful. No markdown, no emojis.",
    "system_style_funny":  "You are Jarvis with a dry, witty sense of humor. Be friendly and lightly playful. ≤2 sentences. No profanity.",

    "max_chars": 300,
    "console_banner": True,
}

# ------------------------------------------------------------
# Logging (rotating)
# ------------------------------------------------------------
logger = logging.getLogger("jarvis")
logger.setLevel(logging.INFO)
rot = logging.handlers.RotatingFileHandler(CONFIG["log_file"], maxBytes=1_000_000, backupCount=3, encoding="utf-8")
fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
rot.setFormatter(fmt)
logger.addHandler(rot)

# ------------------------------------------------------------
# Console spinner
# ------------------------------------------------------------
class Spinner:
    def __init__(self, message: str = "Thinking…", interval: float = 0.1):
        self.msg = message
        self.interval = interval
        self._stop = threading.Event()
        self._t = None
        self.frames = ["⠋","⠙","⠹","⠸","⠼","⠴","⠦","⠧","⠇","⠏"]
    def __enter__(self):
        def run():
            i = 0
            while not self._stop.is_set():
                sys.stdout.write(f"\r{self.frames[i%len(self.frames)]} {self.msg}   ")
                sys.stdout.flush()
                i += 1
                time.sleep(self.interval)
        self._t = threading.Thread(target=run, daemon=True); self._t.start()
        return self
    def __exit__(self, exc_type, exc, tb):
        self._stop.set()
        if self._t: self._t.join(timeout=0.2)
        sys.stdout.write("\r" + " " * 80 + "\r"); sys.stdout.flush()

# ------------------------------------------------------------
# Memory (turns + profile + search cache + reminders)
# ------------------------------------------------------------
class Memory:
    def __init__(self, path: str):
        self.path = path
        self.data: Dict[str, Any] = {
            "turns": [], "search_cache": {}, "reminders": [], "doc_sessions": {}, "profile": {}
        }
        self._load()
    def _load(self):
        try:
            if os.path.exists(self.path):
                with open(self.path, "r", encoding="utf-8") as f:
                    self.data = json.load(f)
            if not isinstance(self.data, dict):
                raise ValueError("memory malformed")
        except Exception:
            self.data = {"turns": [], "search_cache": {}, "reminders": [], "doc_sessions": {}, "profile": {}}
            self._save()
        self.data.setdefault("turns", [])
        self.data.setdefault("search_cache", {})
        self.data.setdefault("reminders", [])
        self.data.setdefault("doc_sessions", {})
        self.data.setdefault("profile", {})
    def _save(self):
        try:
            with open(self.path, "w", encoding="utf-8") as f:
                json.dump(self.data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.warning("Memory save failed: %s", e)
    def log(self, inp: str, out: str, source: str):
        self.data["turns"].append({"ts": datetime.now().isoformat(timespec="seconds"), "input": inp, "reply": out, "source": source})
        self._save()
    def cache_get(self, key: str) -> Optional[Dict[str, Any]]:
        return self.data["search_cache"].get(key.strip().lower())
    def cache_set(self, key: str, value: Dict[str, Any]):
        self.data["search_cache"][key.strip().lower()] = value
        self._save()
    def add_reminder(self, text: str, due: datetime, repeat: Optional[str] = None) -> str:
        rid = f"r{int(time.time()*1000)}"
        self.data["reminders"].append({"id": rid, "text": text, "due": due.isoformat(), "repeat": repeat, "enabled": True})
        self._save(); return rid
    def list_reminders(self) -> List[Dict[str, Any]]: return list(self.data["reminders"])
    def set_reminder_enabled(self, rid: str, enabled: bool) -> bool:
        for r in self.data["reminders"]:
            if r["id"] == rid: r["enabled"] = enabled; self._save(); return True
        return False
    def doc_set(self, sid: str, meta: Dict[str, Any]): self.data["doc_sessions"][sid] = meta; self._save()
    def doc_get(self, sid: str) -> Optional[Dict[str, Any]]: return self.data["doc_sessions"].get(sid)
    def doc_list(self) -> List[Tuple[str, Dict[str, Any]]]: return sorted(self.data["doc_sessions"].items())
    def set_display_name(self, name: str):
        self.data["profile"]["display_name"] = name.strip(); self._save()
    def get_display_name(self) -> Optional[str]:
        n = self.data.get("profile", {}).get("display_name", "").strip()
        return n or None

# ------------------------------------------------------------
# Notes vault
# ------------------------------------------------------------
class Notes:
    def __init__(self, path: str):
        self.path = path
        self.data = {"notes": []}
        self._load()
    def _load(self):
        try:
            if os.path.exists(self.path):
                with open(self.path, "r", encoding="utf-8") as f:
                    self.data = json.load(f)
            if not isinstance(self.data, dict) or "notes" not in self.data:
                raise ValueError("notes malformed")
        except Exception:
            self.data = {"notes": []}; self._save()
    def _save(self):
        try:
            with open(self.path, "w", encoding="utf-8") as f:
                json.dump(self.data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.warning("Notes save failed: %s", e)
    def add(self, text: str, tags: List[str]) -> str:
        nid = f"n{int(time.time()*1000)}"
        self.data["notes"].append({"id": nid, "ts": datetime.now().isoformat(timespec="seconds"), "tags": tags, "text": text})
        self._save(); return nid
    def list(self) -> List[Dict[str, Any]]: return list(self.data["notes"])
    def search(self, query: str) -> List[Dict[str, Any]]:
        q = query.lower(); res=[]
        for n in self.data["notes"]:
            hay = " ".join(n.get("tags", [])) + " " + n.get("text", "")
            if q in hay.lower(): res.append(n)
        return res[:30]
    def summary(self) -> str:
        ns = self.data["notes"]
        if not ns: return "No notes."
        topics = {}
        for n in ns:
            for t in n.get("tags", []):
                topics[t] = topics.get(t, 0) + 1
        top = ", ".join(f"{k}({v})" for k,v in sorted(topics.items(), key=lambda x:-x[1])[:8]) or "untagged"
        return f"{len(ns)} notes; top tags: {top}."

# ------------------------------------------------------------
# TTS (Windows SAPI primary) + fallback to pyttsx4
# ------------------------------------------------------------
class Speaker:
    def __init__(self, rate:int, volume:int, voice_index:int):
        self.backend = None
        self.q: "queue.Queue[Optional[str]]" = queue.Queue()
        self.running = True
        self.on_start = None
        self.on_end = None
        self.sapi_voice = None
        self.sapi = None
        self.pytts = None

        self._init_sapi(rate, volume, voice_index) or self._init_pyttsx4(rate, volume, voice_index)
        self.worker = threading.Thread(target=self._loop, daemon=True); self.worker.start()

    def set_hooks(self, on_start, on_end):
        self.on_start = on_start
        self.on_end = on_end

    def _init_sapi(self, rate:int, volume:int, voice_index:int) -> bool:
        if os.name != "nt": return False
        try:
            import comtypes.client as cc
            self.sapi = cc
            self.sapi_voice = cc.CreateObject("SAPI.SpVoice")
            try: self.sapi_voice.Rate = int(rate)
            except Exception: pass
            try: self.sapi_voice.Volume = int(volume)
            except Exception: pass
            try:
                vs = self.sapi_voice.GetVoices()
                if vs and 0 <= voice_index < vs.Count:
                    self.sapi_voice.Voice = vs.Item(voice_index)
                # Try auto‑select British
                for i in range(vs.Count):
                    d = (vs.Item(i).GetDescription() if hasattr(vs.Item(i),"GetDescription") else "").lower()
                    if "english" in d and ("united kingdom" in d or "great britain" in d or "hazel" in d):
                        self.sapi_voice.Voice = vs.Item(i); break
            except Exception: pass
            self.backend = "sapi"
            return True
        except Exception as e:
            logger.warning("SAPI init failed: %s", e)
            return False

    def _init_pyttsx4(self, rate:int, volume:int, voice_index:int) -> bool:
        try:
            import pyttsx4
            self.pytts = pyttsx4.init()
            try: self.pytts.setProperty("rate", 180 if rate==0 else 180+(rate*10))
            except Exception: pass
            try: self.pytts.setProperty("volume", max(0.0, min(1.0, volume/100.0)))
            except Exception: pass
            voices = []
            try: voices = self.pytts.getProperty("voices")
            except Exception: pass
            if voices and 0 <= voice_index < len(voices):
                try: self.pytts.setProperty("voice", voices[voice_index].id)
                except Exception: pass
            self.backend = "pyttsx4"
            return True
        except Exception as e:
            logger.warning("pyttsx4 init failed: %s", e)
            self.backend = None
            return False

    def _loop(self):
        while self.running:
            text = self.q.get()
            if text is None: break
            if not text.strip(): continue
            try:
                if self.on_start:
                    try: self.on_start()
                    except Exception: pass
                print("🔊 speaking…")
                if self.backend == "sapi":
                    self.sapi_voice.Speak(text, 0)
                elif self.backend == "pyttsx4":
                    try:
                        self.pytts.say(text); self.pytts.runAndWait()
                    except Exception as e:
                        logger.warning("pyttsx4 speak warn: %s", e)
            except Exception as e:
                logger.warning("TTS say warn: %s", e)
            finally:
                if self.on_end:
                    try: self.on_end()
                    except Exception: pass

    def say(self, text:str):
        self.q.put(text)

    def stop(self):
        self.running = False
        self.q.put(None)

# ------------------------------------------------------------
# Helpers
# ------------------------------------------------------------
def resample_to_16k(x: np.ndarray, sr: int) -> np.ndarray:
    x = np.asarray(x, dtype=np.float32).squeeze()
    if sr == 16000: return x
    if len(x) < 2: return np.zeros((0,), dtype=np.float32)
    duration = len(x) / float(sr)
    new_len = int(round(duration * 16000))
    t_old = np.linspace(0.0, duration, num=len(x), endpoint=False, dtype=np.float64)
    t_new = np.linspace(0.0, duration, num=new_len, endpoint=False, dtype=np.float64)
    y = np.interp(t_new, t_old, x.astype(np.float64))
    return y.astype(np.float32)

def float_to_pcm16(arr: np.ndarray) -> bytes:
    a = np.clip(arr, -1.0, 1.0)
    return (a * 32767.0).astype(np.int16).tobytes()

def frame_bytes_20ms(arr16: np.ndarray) -> List[bytes]:
    # 20ms @ 16kHz = 320 samples
    hop = 320
    out = []
    for i in range(0, len(arr16) - hop + 1, hop):
        chunk = arr16[i:i+hop]
        out.append(float_to_pcm16(chunk))
    return out

# ------------------------------------------------------------
# STT Listener
# ------------------------------------------------------------
class Listener:
    def __init__(self, sample_rate:int, wake_enabled:bool, wake_words:List[str]):
        self.rate = sample_rate
        self.device: Optional[int] = None
        self.device_name = "default"

        self.wake_enabled = wake_enabled
        self.wake_words = [w.lower() for w in wake_words]

        self.model_name = CONFIG.get("whisper_model","medium.en")
        self.model_lock = threading.Lock()
        self.model = WhisperModel(self.model_name, device="auto", compute_type="int8")

        self.running = False
        self.thread = None
        self.out_q: "queue.Queue[Tuple[str,str]]" = queue.Queue()

        self.silence_ms = CONFIG.get("vad_silence_ms", 300)
        self.pre_ms = CONFIG.get("vad_pre_speech_ms", 450)
        self.max_utt_s = CONFIG.get("vad_max_utterance_s", 12)
        self.min_utt_s = CONFIG.get("vad_min_utterance_s", 0.30)

        self.energy_thresh = 0.01
        self.pause_event = threading.Event()

        # Optional WebRTC VAD
        self.use_webrtc = CONFIG.get("webrtc_vad", True) and WEBRTC_OK
        if self.use_webrtc:
            self.vad = webrtcvad.Vad()
            try:
                self.vad.set_mode(int(CONFIG.get("webrtc_vad_aggressiveness", 2)))
            except Exception:
                self.vad.set_mode(2)

    def set_model(self, model_name: str) -> str:
        with self.model_lock:
            try:
                self.model = WhisperModel(model_name, device="auto", compute_type="int8")
                self.model_name = model_name
                CONFIG["whisper_model"] = model_name
                return f"STT model set to {model_name}."
            except Exception as e:
                return f"Model load failed: {e}"

    def _prepare_device(self):
        try:
            devs = sd.query_devices()
            defaults = sd.default.device
            forced = CONFIG.get("input_device")
            if forced is not None:
                d = devs[forced]; name = d.get("name","?"); rate = int(d.get("default_samplerate") or 16000)
                self.device = forced; self.device_name = name; self.rate = rate
            else:
                if isinstance(defaults,(list,tuple)) and defaults and defaults[0] is not None and defaults[0] >= 0:
                    idx = int(defaults[0]); d = devs[idx]; name = d.get("name","?"); rate = int(d.get("default_samplerate") or 16000)
                    self.device = idx; self.device_name = name; self.rate = rate
                else:
                    for i,d in enumerate(devs):
                        if d.get("max_input_channels",0) > 0:
                            name=d.get("name","?"); rate=int(d.get("default_samplerate") or 16000)
                            self.device = i; self.device_name = name; self.rate = rate; break
        except Exception as e:
            logger.warning("pick_input_device warn: %s", e)
        if not self.rate: self.rate = CONFIG.get("sample_rate",16000)
        print(f"🎤 Mic: {self.device_name} (idx={self.device if self.device is not None else 'default'}), {self.rate} Hz")

    def calibrate(self, seconds:float=0.6):
        try:
            buf = sd.rec(int(seconds * self.rate), samplerate=self.rate, channels=1, dtype="float32", device=self.device)
            sd.wait()
            # median is more robust to spikes
            m = float(np.median(np.abs(buf)))
            self.energy_thresh = max(0.0025, m * 2.0)
            logger.info("Mic calibrated: median=%.6f, energy_thresh=%.4f", m, self.energy_thresh)
        except Exception as e:
            print(f"⚠️ Mic calibration failed; voice disabled (type to interact). Error: {e}")
            logger.warning("Mic calibrate warn: %s", e)

    def pause(self): self.pause_event.set()
    def resume(self): self.pause_event.clear()

    def _listen_once(self) -> str:
        chunk_s = 0.05
        chunk_n = int(self.rate * chunk_s)
        pre_frames = int(self.pre_ms / 1000 / chunk_s)
        silence_frames_needed = int(self.silence_ms / 1000 / chunk_s)
        max_chunks = int(self.max_utt_s / chunk_s)

        ring: List[np.ndarray] = []
        speech: List[np.ndarray] = []
        in_speech = False
        silence_count = 0
        chunks = 0

        try:
            with sd.InputStream(samplerate=self.rate, channels=1, dtype="float32", blocksize=0, device=self.device) as stream:
                while True:
                    if self.pause_event.is_set():
                        time.sleep(0.05); continue
                    audio, _ = stream.read(chunk_n)
                    chunks += 1
                    if chunks > max_chunks: break

                    # energy VAD
                    energy = float(np.mean(np.abs(audio)))
                    energy_voice = energy > self.energy_thresh

                    # WebRTC VAD (on resampled 16k 20ms frames)
                    vad_voice = False
                    if self.use_webrtc:
                        arr16 = resample_to_16k(audio.squeeze(), self.rate)
                        frames = frame_bytes_20ms(arr16)
                        for fb in frames:
                            if self.vad.is_speech(fb, 16000):
                                vad_voice = True; break

                    is_voice = energy_voice or vad_voice

                    if not in_speech:
                        ring.append(audio.copy())
                        if len(ring) > pre_frames: ring.pop(0)
                        if is_voice:
                            in_speech = True; speech.extend(ring); ring.clear(); silence_count = 0
                    else:
                        speech.append(audio.copy())
                        if is_voice:
                            silence_count = 0
                        else:
                            silence_count += 1
                            if silence_count >= silence_frames_needed: break
        except Exception as e:
            logger.warning("InputStream warn: %s", e)
            print(f"⚠️ Mic open failed; voice disabled (type to interact). Error: {e}")
            return ""

        if not speech: return ""

        arr = np.concatenate(speech, axis=0).astype(np.float32).squeeze()
        dur_raw = len(arr) / float(self.rate)
        if dur_raw < self.min_utt_s: return ""

        arr16 = resample_to_16k(arr, self.rate)
        if len(arr16) < int(16000 * 0.30): return ""

        try:
            lang = CONFIG.get("whisper_language","en")
            lang = None if (not lang or lang == "auto") else lang
            with self.model_lock:
                segments, _ = self.model.transcribe(
                    arr16,
                    language=lang,
                    beam_size=CONFIG.get("whisper_beam_size", 5),
                    temperature=CONFIG.get("whisper_temperature", 0.0),
                    vad_filter=False,
                    without_timestamps=True,
                    condition_on_previous_text=False,
                    initial_prompt="Jarvis; open, google, youtube, gmail, calendar, drive, github; note; remind; call me; solve; calc; hello; yes; okay."
                )
            text = " ".join(seg.text for seg in segments).strip()
            return text
        except Exception as e:
            logger.warning("Transcribe warn: %s", e)
            return ""

    def _loop(self):
        self._prepare_device()
        self.calibrate(0.6)
        while self.running:
            try:
                if self.pause_event.is_set():
                    time.sleep(0.05); continue
                text = self._listen_once()
                if not text: continue
                if self.wake_enabled:
                    if self._has_wake(text):
                        for w in self.wake_words:
                            idx = text.lower().find(w)
                            if idx != -1:
                                text = text[idx + len(w):].strip(" ,.!?") or "ready"; break
                        self.out_q.put(("voice", text))
                else:
                    self.out_q.put(("voice", text))
            except Exception as e:
                logger.warning("Listener loop warn: %s", e)
                print("⚠️ Listener error — recovering…")
                time.sleep(1.0)
                self._prepare_device(); self.calibrate(0.5)

    def _has_wake(self, text:str) -> bool:
        t = text.lower()
        return any(w in t for w in self.wake_words)

    def start(self):
        if self.running: return
        self.running = True
        self.thread = threading.Thread(target=self._loop, daemon=True); self.thread.start()
    def stop(self): self.running = False

# ------------------------------------------------------------
# Typist
# ------------------------------------------------------------
class Typist:
    def __init__(self):
        self.thread=None; self.running=False
        self.out_q: "queue.Queue[Tuple[str,str]]" = queue.Queue()
    def _loop(self):
        while self.running:
            try:
                s=input("")
                if s.strip():
                    self.out_q.put(("text", s.strip()))
            except EOFError:
                time.sleep(0.1)
            except Exception:
                time.sleep(0.1)
    def start(self):
        if self.running: return
        self.running=True; self.thread=threading.Thread(target=self._loop, daemon=True); self.thread.start()
    def stop(self): self.running=False

# ------------------------------------------------------------
# Command Registry
# ------------------------------------------------------------
class CommandRegistry:
    def __init__(self): self.commands: Dict[str, Dict[str, Any]] = {}
    def register(self, name:str, func, help_text:str, aliases:List[str]=None):
        self.commands[name] = {"func": func, "help": help_text, "aliases": aliases or []}
    def resolve(self, text:str) -> Optional[Tuple[str, Dict[str, Any]]]:
        t = text.strip().lower()
        for k, meta in self.commands.items():
            for nm in [k] + meta["aliases"]:
                if t == nm or t.startswith(nm + " "): return k, meta
        return None
    def help_markdown(self) -> str:
        lines = []
        for k, meta in sorted(self.commands.items()):
            al = (", ".join(meta["aliases"])) if meta["aliases"] else ""
            lines.append(f"- {k} {f'({al})' if al else ''}: {meta['help']}")
        return "\n".join(lines)

# ------------------------------------------------------------
# Tools (web, files, pdf, etc.)
# ------------------------------------------------------------
class Tools:
    def __init__(self, mem:Memory): self.mem = mem
    def web_lookup(self, query:str) -> Dict[str, Any]:
        key = CONFIG["serpapi_key"]
        if not key: return {"text":"Web search disabled (no SERPAPI_KEY).", "confidence":0.0}
        cached = self.mem.cache_get(query)
        if cached: return cached
        try:
            r = requests.get("https://serpapi.com/search", params={"engine":"google","q":query,"api_key":key}, timeout=10)
            js = r.json()
            text = None; conf = 0.2
            ab = js.get("answer_box") or {}; kg = js.get("knowledge_graph") or {}; org = (js.get("organic_results") or [{}])
            if ab.get("answer"): text = ab["answer"]; conf = 0.95
            elif ab.get("snippet"): text = ab["snippet"]; conf = 0.9
            elif kg.get("description"): text = kg["description"]; conf = 0.8
            elif org and org[0].get("snippet"): text = org[0]["snippet"]; conf = 0.6
            else: text = "No clear answer online."; conf = 0.3
            out = {"text": text, "confidence": float(conf)}
            self.mem.cache_set(query, out); return out
        except Exception:
            return {"text":"Search error.", "confidence":0.0}
    def web_answer(self, query:str) -> str: return self.web_lookup(query)["text"]
    def fetch_json(self, url:str) -> str:
        try:
            r = requests.get(url, timeout=10)
            if r.ok:
                js = r.json(); s = json.dumps(js, indent=2)[:1000]
                return s + ("..." if len(s)==1000 else "")
            return f"HTTP {r.status_code}"
        except Exception: return "Fetch error."
    def weather(self, place:str) -> str:
        try:
            g = requests.get("https://geocoding-api.open-meteo.com/v1/search", params={"name":place, "count":1}, timeout=10).json()
            if not g.get("results"): return "Place not found."
            lat = g["results"][0]["latitude"]; lon = g["results"][0]["longitude"]
            w = requests.get("https://api.open-meteo.com/v1/forecast", params={"latitude": lat, "longitude": lon, "current_weather":"true"}, timeout=10).json()
            cw = w.get("current_weather") or {}; temp = cw.get("temperature"); wind = cw.get("windspeed")
            if temp is None: return "Weather unavailable."
            return f"{place}: {temp}°C, wind {wind} km/h."
        except Exception: return "Weather error."
    def list_files(self, path:str=".") -> str:
        try:
            entries = os.listdir(path); parts=[]
            for e in entries[:200]:
                p = os.path.join(path, e); tag = "DIR" if os.path.isdir(p) else "FILE"; parts.append(f"{tag}\t{e}")
            return "\n".join(parts) if parts else "Empty."
        except Exception: return "Path error."
    def read_file(self, path:str) -> str:
        try:
            with open(path,"r",encoding="utf-8",errors="ignore") as f: s=f.read()
            return s[:1500]+("..." if len(s)>1500 else "")
        except Exception: return "Read error."
    def write_file(self, path:str, content:str) -> str:
        try:
            with open(path,"w",encoding="utf-8") as f: f.write(content); return "Saved."
        except Exception: return "Write error."
    def copy(self, src:str, dst:str)->str:
        try:
            if os.path.isdir(src):
                if os.path.isdir(dst): dst = os.path.join(dst, os.path.basename(src))
                shutil.copytree(src, dst)
            else: shutil.copy2(src, dst)
            return "Copied."
        except Exception: return "Copy error."
    def move(self, src:str, dst:str)->str:
        try: shutil.move(src, dst); return "Moved."
        except Exception: return "Move error."
    def rename(self, src:str, dst:str)->str:
        try: os.rename(src, dst); return "Renamed."
        except Exception: return "Rename error."
    def delete(self, path:str)->str:
        try:
            if os.path.isdir(path): shutil.rmtree(path)
            else: os.remove(path)
            return "Deleted."
        except Exception: return "Delete error."
    def zip_create(self, src:str, zip_path:str)->str:
        try:
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
                if os.path.isdir(src):
                    for root, _, files in os.walk(src):
                        for f in files:
                            ab = os.path.join(root, f); rel = os.path.relpath(ab, src); z.write(ab, rel)
                else: z.write(src, os.path.basename(src))
            return "Zipped."
        except Exception: return "Zip error."
    def unzip(self, zip_path:str, out_dir:str)->str:
        try:
            with zipfile.ZipFile(zip_path, "r") as z: z.extractall(out_dir); return "Unzipped."
        except Exception: return "Unzip error."
    def text_to_pdf(self, text:str, out_path:str)->str:
        try:
            c = canvas.Canvas(out_path, pagesize=LETTER)
            width, height = LETTER; left = 1*inch; top = height - 1*inch; lh = 14; y=top
            for line in text.splitlines():
                if y < 1*inch: c.showPage(); y=top
                c.drawString(left, y, line[:110]); y -= lh
            c.save(); return f"PDF created: {out_path}"
        except Exception: return "PDF error."
    def docx_to_pdf(self, docx_path:str, pdf_path:str)->str:
        try:
            if DOCX2PDF_AVAILABLE:
                docx2pdf_convert(docx_path, pdf_path); return f"Converted: {pdf_path}"
        except Exception: pass
        try:
            import comtypes.client as cc
            word = cc.CreateObject('Word.Application'); word.Visible=False
            doc=word.Documents.Open(os.path.abspath(docx_path))
            doc.ExportAsFixedFormat(os.path.abspath(pdf_path), 17)
            doc.Close(False); word.Quit(); return f"Converted: {pdf_path}"
        except Exception: return "DOCX→PDF failed."
    def pdf_to_docx(self, pdf_path:str, docx_path:str)->str:
        try:
            full=[]
            with pdfplumber.open(pdf_path) as pdf:
                for pg in pdf.pages: full.append(pg.extract_text() or "")
            doc = DocxDocument()
            for line in "\n".join(full).splitlines():
                p = doc.add_paragraph(); run = p.add_run(line); run.font.size=Pt(11)
            doc.save(docx_path); return f"Converted (text-only): {docx_path}"
        except Exception: return "PDF→DOCX failed."
    def pdf_merge(self, out_path:str, *pdfs:str)->str:
        try:
            writer = PdfWriter()
            for p in pdfs:
                r = PdfReader(p)
                for pg in r.pages: writer.add_page(pg)
            with open(out_path,"wb") as f: writer.write(f)
            return f"Merged: {out_path}"
        except Exception: return "Merge failed."
    def pdf_split(self, pdf_path:str, out_dir:str)->str:
        try:
            r = PdfReader(pdf_path)
            for i, pg in enumerate(r.pages):
                w = PdfWriter(); w.add_page(pg)
                p = os.path.join(out_dir, f"page_{i+1}.pdf")
                with open(p,"wb") as f: w.write(f)
            return f"Split to {out_dir}"
        except Exception: return "Split failed."
    def pdf_extract(self, pdf_path:str, out_path:str, start:int, end:int)->str:
        try:
            r=PdfReader(pdf_path); w=PdfWriter()
            n=len(r.pages); start=max(1, start); end=min(n, end)
            if start>end: return "Range invalid."
            for i in range(start-1, end): w.add_page(r.pages[i])
            with open(out_path,"wb") as f: w.write(f)
            return f"Extracted pages {start}-{end}."
        except Exception: return "Extract failed."
    def doc_load(self, sid:str, path:str) -> str:
        try:
            text = ""
            if path.lower().endswith(".pdf"):
                with pdfplumber.open(path) as pdf:
                    for pg in pdf.pages: text += (pg.extract_text() or "") + "\n"
            elif path.lower().endswith(".docx"):
                d = DocxDocument(path)
                for p in d.paragraphs: text += p.text + "\n"
            elif path.lower().endswith(".txt"):
                with open(path,"r",encoding="utf-8",errors="ignore") as f: text = f.read()
            else: return "Unsupported."
            text = text[:200000]
            self.mem.doc_set(sid, {"path": os.path.abspath(path), "text": text})
            return f"Doc {sid} loaded."
        except Exception: return "Load failed."
    def doc_ask(self, sid:str, question:str, llm) -> str:
        sess = self.mem.doc_get(sid)
        if not sess: return "No session."
        context = sess.get("text","")[:8000]
        prompt = f"Context:\n{context}\n\nQuestion: {question}\nAnswer briefly:"
        return llm.ask(prompt)
    def clip_get(self)->str:
        try: return pyperclip.paste()[:1000]
        except Exception: return "Clipboard read error."
    def clip_set(self, text:str)->str:
        try: pyperclip.copy(text); return "Clipboard updated."
        except Exception: return "Clipboard write error."
    def screenshot(self, out_path:str)->str:
        try:
            import pyautogui
            im = pyautogui.screenshot(); im.save(out_path)
            return f"Screenshot saved: {out_path}"
        except Exception: return "Screenshot failed."

# ------------------------------------------------------------
# Calculator / Unit Conversion
# ------------------------------------------------------------
class SafeCalc:
    ALLOWED = {k:getattr(math,k) for k in dir(math) if not k.startswith("_")}
    ALLOWED.update({"abs":abs, "round":round, "min":min, "max":max})
    UNITS = {
        "m":1.0, "km":1000.0, "cm":0.01, "mm":0.001, "mi":1609.344, "ft":0.3048, "in":0.0254, "yd":0.9144,
        "kg":1.0, "g":0.001, "lb":0.45359237, "oz":0.0283495231,
        "c":1.0, "f":"_tempf", "k":"_tempk"
    }
    def calc(self, expr:str)->str:
        t = expr.strip()
        if " to " in t.lower(): return self._convert(t)
        try:
            code = compile(t, "<calc>", "eval")
            for n in code.co_names:
                if n not in self.ALLOWED: return "Invalid."
            val = eval(code, {"__builtins__":{}}, self.ALLOWED)
            return str(val)
        except Exception: return "Invalid."
    def _convert(self, s:str)->str:
        try:
            parts = re.split(r"\s+to\s+", s, flags=re.I)
            left, unit_to = parts[0].strip(), parts[1].strip()
            m = re.match(r"^\s*([-+]?\d+(\.\d+)?)\s*([A-Za-z]+)\s*$", left)
            if not m: return "Invalid."
            val = float(m.group(1)); unit_from = m.group(3).lower()
            unit_to = unit_to.lower()
            if unit_from in ("c","f","k") or unit_to in ("c","f","k"):
                return self._convert_temp(val, unit_from, unit_to)
            if unit_from not in self.UNITS or unit_to not in self.UNITS: return "Unit?"
            base = val * self.UNITS[unit_from]; out = base / self.UNITS[unit_to]
            return f"{out:g} {unit_to}"
        except Exception: return "Invalid."
    def _convert_temp(self, v:float, f:str, t:str)->str:
        def c2f(x): return x*9/5+32
        def f2c(x): return (x-32)*5/9
        def c2k(x): return x+273.15
        def k2c(x): return x-273.15
        f=f.lower(); t=t.lower()
        if f==t: return f"{v:g} {t}"
        if f=="c": c=v
        elif f=="f": c=f2c(v)
        elif f=="k": c=k2c(v)
        else: return "Unit?"
        if t=="c": return f"{c:g} c"
        if t=="f": return f"{c2f(c):g} f"
        if t=="k": return f"{c2k(c):g} k"
        return "Unit?"

# ------------------------------------------------------------
# Math engine (SymPy)
# ------------------------------------------------------------
class MathEngine:
    def is_mathy(self, text:str) -> bool:
        t = text.lower()
        return any(k in t for k in [
            "mod", "φ", "phi", "gcd", "lcm", "prime", "factor", "solve",
            "rsa", "inverse", "modular", "congruent", "matrix", "det",
            "crt", "totient", "euler", "log", "sqrt", "^", "**", "=", "≡"
        ]) or bool(re.search(r"\d", t))
    def solve_expr(self, expr:str) -> str:
        try:
            expr = expr.strip().lstrip("solve").strip()
            if "≡" in expr: expr = expr.replace("≡", "==")
            if "φ" in expr: expr = expr.replace("φ", "phi")
            loc = {"phi": sp.totient, "gcd": sp.gcd, "lcm": sp.lcm, "sqrt": sp.sqrt, "Matrix": sp.Matrix}
            res = sp.sympify(expr, locals=loc)
            if isinstance(res, sp.Equality):
                syms = list(res.free_symbols) or [sp.Symbol("x")]
                sol = sp.solve(res, *syms, dict=True)
                return str(sol) if sol else "No solution."
            val = sp.simplify(res)
            return str(val)
        except Exception:
            try:
                x=sp.Symbol("x")
                val = eval(expr, {"__builtins__":{}}, {"sp":sp, "phi":sp.totient, "gcd":sp.gcd, "lcm":sp.lcm, "sqrt":sp.sqrt, "Matrix":sp.Matrix})
                return str(val)
            except Exception:
                return "Invalid."
    def factor(self, n:int) -> str:
        try:
            n = int(str(n).strip()); fac = sp.factorint(n)
            parts = [f"{p}^{e}" if e>1 else f"{p}" for p,e in fac.items()]
            return " * ".join(parts) if parts else "1"
        except Exception: return "Invalid."
    def mod_inv(self, a:int, m:int) -> str:
        try:
            a=int(a); m=int(m); inv = sp.mod_inverse(a, m); return str(inv)
        except Exception: return "No inverse."
    def crt(self, mods:List[int], rems:List[int]) -> str:
        try:
            x, mod = sp.ntheory.modular.crt(mods, rems)
            return f"x ≡ {int(x)} (mod {int(mod)})"
        except Exception: return "CRT error."
    def rsa_from_n_e(self, n:int, e:int) -> str:
        try:
            n=int(n); e=int(e); fac = sp.factorint(n)
            if len(fac)!=2: return "n is not semiprime."
            (p,_),(q,_) = list(fac.items()); phi = (p-1)*(q-1)
            try: d = sp.mod_inverse(e, phi)
            except Exception: return "No inverse."
            return f"p={p}, q={q}, phi={phi}, d={d}"
        except Exception: return "Invalid."

# ------------------------------------------------------------
# Local LLM (Ollama) with conversation history
# ------------------------------------------------------------
class LocalLLM:
    def __init__(self, api:str, model:str, timeout:int, system_prompt:str, reply_lang:str="en"):
        self.api = api; self.model = model; self.timeout = timeout
        self.system = system_prompt
        self.reply_lang = reply_lang

    def ask(self, user_text:str, history:Optional[List[Dict[str,str]]]=None, reply_lang:Optional[str]=None, user_name:Optional[str]=None) -> str:
        lang = (reply_lang or self.reply_lang or "en").strip()
        name_hint = f"When addressing the user, call them '{user_name}'." if user_name else ""
        lines = [self.system, f"Respond in {lang}.", name_hint, ""]
        if history:
            for turn in history[-8:]:
                role = turn.get("role","user").capitalize()
                txt  = turn.get("text","").strip()
                if not txt: continue
                lines.append(f"{role}: {txt}")
        lines.append(f"User: {user_text}")
        lines.append("Jarvis:")
        prompt = "\n".join([x for x in lines if x is not None])

        try:
            r = requests.post(self.api, json={"model": self.model, "prompt": prompt, "stream": False}, timeout=self.timeout)
            reply = ""
            try: reply = r.json().get("response","").strip()
            except Exception: reply = (r.text or "").strip()
            flat = " ".join(line.strip() for line in (reply or "Ready.").splitlines() if line.strip())
            if not flat or flat.lower() == "ready.":
                # deterministic fallbacks for common utterances
                u = user_text.strip().lower()
                if "can you hear me" in u or "hear me" in u: return "Loud and clear."
                if u in ("who are you", "who r u"): return "I'm Jarvis—your brisk, British‑leaning sidekick."
                if u in ("hi", "hello", "hey"): return "Hello. How can I help?"
                return "Got it."
            if len(flat) > CONFIG["max_chars"]:
                flat = flat[:CONFIG["max_chars"]].rsplit(" ",1)[0] + "..."
            return flat.replace("**","").replace("*","").strip()
        except Exception as e:
            logger.warning("LLM error: %s", e); return "Okay."

# ------------------------------------------------------------
# Scheduler / Reminders
# ------------------------------------------------------------
class ReminderScheduler:
    def __init__(self, memory:Memory, speaker:Speaker):
        self.mem = memory; self.speaker = speaker
        self.thread = None; self.running = False
    def _loop(self):
        while self.running:
            try:
                now = datetime.now(); changed = False
                for r in self.mem.list_reminders():
                    if not r.get("enabled", True): continue
                    due = datetime.fromisoformat(r["due"])
                    if now >= due:
                        self.speaker.say(f"Reminder: {r['text']}")
                        rep = r.get("repeat")
                        if rep == "daily":
                            due = due + timedelta(days=1); r["due"] = due.isoformat()
                        elif rep == "weekly":
                            due = due + timedelta(weeks=1); r["due"] = due.isoformat()
                        else:
                            r["enabled"] = False
                        changed = True
                if changed: self.mem._save()
                time.sleep(1)
            except Exception as e:
                logger.warning("Reminder loop warn: %s", e); time.sleep(2)
    def start(self):
        if self.running: return
        self.running = True; self.thread = threading.Thread(target=self._loop, daemon=True); self.thread.start()
    def stop(self): self.running = False

# ------------------------------------------------------------
# Orchestrator
# ------------------------------------------------------------
class JarvisPrime:
    def __init__(self):
        self.mem = Memory(CONFIG["memory_file"])
        self.notes = Notes(CONFIG["notes_file"])
        self.speaker = Speaker(CONFIG["voice_rate"], CONFIG["voice_volume"], CONFIG["voice_index"])
        self.listener = Listener(CONFIG["sample_rate"], CONFIG["wake_word_enabled"], CONFIG["wake_words"])
        self.speaker.set_hooks(self.listener.pause, self.listener.resume)

        self.typist = Typist()
        self.tools = Tools(self.mem)
        self.calc = SafeCalc()
        self.math = MathEngine()

        self.history: Deque[Dict[str,str]] = deque(maxlen=16)
        self._bootstrap_history_from_memory(max_pairs=6)

        self.user_name = self.mem.get_display_name()

        self.llm_concise = LocalLLM(CONFIG["ollama_api"], CONFIG["ollama_model"], CONFIG["ollama_timeout_s"], CONFIG["system_style_concise"], CONFIG["reply_language"])
        self.llm_funny   = LocalLLM(CONFIG["ollama_api"], CONFIG["ollama_model"], CONFIG["ollama_timeout_s"], CONFIG["system_style_funny"],   CONFIG["reply_language"])

        self.scheduler = ReminderScheduler(self.mem, self.speaker)
        self.in_q: "queue.Queue[Tuple[str,str]]" = queue.Queue()
        self.running = False

        self.commands = CommandRegistry()
        self._register_commands()

    def _bootstrap_history_from_memory(self, max_pairs:int=6):
        turns = self.mem.data.get("turns", [])
        if not turns: return
        for t in turns[-max_pairs:]:
            u = t.get("input","").strip()
            a = t.get("reply","").strip()
            if u: self.history.append({"role":"user","text":u})
            if a: self.history.append({"role":"assistant","text":a})

    # ------------------ Commands ------------------
    def _register_commands(self):
        R = self.commands.register
        # help / system
        R("help", self.cmd_help, "Show commands.", aliases=["?", "commands"])
        R("system", self.cmd_system, "Show CPU and RAM.")
        R("exit", self.cmd_exit, "Exit.", aliases=["quit","bye"])

        # web / fetch / weather
        R("search", self.cmd_search, "Web search. Usage: search <query>")
        R("fetch", self.cmd_fetch, "GET JSON. Usage: fetch <url>")
        R("weather", self.cmd_weather, "Weather. Usage: weather <city>")

        # files
        R("files", self.cmd_files, "List. Usage: files [path]")
        R("read", self.cmd_read, "Read preview. Usage: read <path>")
        R("write", self.cmd_write, "Write. Usage: write <path> :: <content>")
        R("copy", self.cmd_copy, "Copy. Usage: copy <src> :: <dst>")
        R("move", self.cmd_move, "Move. Usage: move <src> :: <dst>")
        R("rename", self.cmd_rename, "Rename. Usage: rename <src> :: <dst>")
        R("delete", self.cmd_delete, "Delete. Usage: delete <path>")
        R("zip", self.cmd_zip, "Zip. Usage: zip <src> :: <zip_path>")
        R("unzip", self.cmd_unzip, "Unzip. Usage: unzip <zip> :: <out_dir>")

        # PDFs/DOCX
        R("makepdf", self.cmd_makepdf, "Text→PDF. Usage: makepdf <out.pdf> :: <text>")
        R("docx2pdf", self.cmd_docx2pdf, "DOCX→PDF. Usage: docx2pdf <in.docx> :: <out.pdf>")
        R("pdf2docx", self.cmd_pdf2docx, "PDF→DOCX (text). Usage: pdf2docx <in.pdf> :: <out.docx>")
        R("pdfmerge", self.cmd_pdfmerge, "Merge. Usage: pdfmerge <out.pdf> :: <a.pdf>;<b.pdf>[;...]")
        R("pdfsplit", self.cmd_pdfsplit, "Split. Usage: pdfsplit <in.pdf> :: <out_dir>")
        R("pdfextract", self.cmd_pdfextract, "Extract. Usage: pdfextract <in.pdf> :: <out.pdf> :: <start>-<end>")

        # doc chat
        R("docload", self.cmd_docload, "Load doc. Usage: docload <session_id> :: <path>")
        R("docs", self.cmd_docs, "List doc sessions.")
        R("docask", self.cmd_docask, "Ask doc. Usage: docask <session_id> :: <question>")

        # clipboard / screenshot / open
        R("clipget", self.cmd_clipget, "Clipboard read.")
        R("clipset", self.cmd_clipset, "Clipboard write. Usage: clipset :: <text>")
        R("screenshot", self.cmd_screenshot, "Screenshot. Usage: screenshot <out.png>")
        R("open", self.cmd_open, "Open url/app. Usage: open <url|app>")

        # notes
        R("note", self.cmd_note, "Add note. Usage: note [#tag #tag2] :: <text>")
        R("notes", self.cmd_notes, "List or search. Usage: notes [text]")
        R("notesum", self.cmd_notesum, "Notes summary.")

        # reminders
        R("remind", self.cmd_remind, "Set reminder. Usage: remind <YYYY-MM-DD HH:MM> :: <text>")
        R("reminders", self.cmd_reminders, "List reminders.")
        R("remindon", self.cmd_remindon, "Enable. Usage: remindon <id>")
        R("remindoff", self.cmd_remindoff, "Disable. Usage: remindoff <id>")

        # audio / tts
        R("audio", self.cmd_audio, "Audio: list/select/tune. Usage: audio [list|in <index>|autotune]")
        R("rate", self.cmd_rate, "Set TTS rate (SAPI -10..+10). Usage: rate <int>")
        R("volume", self.cmd_volume, "Set TTS volume (0..100). Usage: volume <int>")
        R("say", self.cmd_say, "Speak a sentence. Usage: say <text>")

        # calc / math
        R("calc", self.cmd_calc, "Calculator / convert. Usage: calc <expr> or '5 km to m'")
        R("solve", self.cmd_solve, "SymPy solve/simplify. Usage: solve <expr>")
        R("factor", self.cmd_factor, "Prime factorization. Usage: factor <n>")
        R("modinv", self.cmd_modinv, "Mod inverse. Usage: modinv <a> :: <m>")
        R("crt", self.cmd_crt, "CRT. Usage: crt <m1,m2,...> :: <r1,r2,...>")
        R("rsa", self.cmd_rsa, "RSA helper. Usage: rsa <n> :: <e>")

        # language controls
        R("lang", self.cmd_lang, "Language: show/set. Usage: lang show | lang stt <code|auto> | lang reply <code>")

        # profile
        R("callme", self.cmd_callme, "Set your display name. Usage: callme <name>")
        R("whoami", self.cmd_whoami, "Show your display name.")

        # STT controls
        R("stt", self.cmd_stt, "STT controls. Usage: stt quality <best|high|balanced|fast> | stt model <name> | stt reload")

        # debug
        R("debug", self.cmd_debug, "Debug tools. Usage: debug stt <seconds>")

    # ------------------ Command handlers ------------------
    def cmd_help(self, arg:str) -> str: return self.commands.help_markdown()
    def cmd_system(self, arg:str)->str:
        cpu = psutil.cpu_percent(interval=0.3); mem = psutil.virtual_memory().percent
        now = datetime.now().strftime("%I:%M %p, %A, %B %d")
        return f"{now}. CPU {cpu}%, RAM {mem}%."
    def cmd_exit(self, arg:str)->str: self.running=False; return "Goodbye."

    def cmd_search(self, arg:str)->str:
        if not arg.strip(): return "Usage: search <query>"
        with Spinner("Searching…"):
            res = self.tools.web_lookup(arg.strip())
        return res["text"]
    def cmd_fetch(self, arg:str)->str:
        if not arg.strip(): return "Usage: fetch <url>"
        with Spinner("Fetching…"):
            return self.tools.fetch_json(arg.strip())
    def cmd_weather(self, arg:str)->str:
        if not arg.strip(): return "Usage: weather <city>"
        with Spinner("Getting weather…"):
            return self.tools.weather(arg.strip())

    def cmd_files(self, arg:str)->str: return self.tools.list_files(arg.strip() or ".")
    def cmd_read(self, arg:str)->str:
        if not arg.strip(): return "Usage: read <path>"
        return self.tools.read_file(arg.strip())
    def cmd_write(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: write <path> :: <content>"
        return self.tools.write_file(parts[0].strip(), parts[1].strip())
    def cmd_copy(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: copy <src> :: <dst>"
        return self.tools.copy(parts[0].strip(), parts[1].strip())
    def cmd_move(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: move <src> :: <dst>"
        return self.tools.move(parts[0].strip(), parts[1].strip())
    def cmd_rename(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: rename <src> :: <dst>"
        return self.tools.rename(parts[0].strip(), parts[1].strip())
    def cmd_delete(self, arg:str)->str:
        if not arg.strip(): return "Usage: delete <path>"
        return self.tools.delete(arg.strip())
    def cmd_zip(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: zip <src> :: <zip_path>"
        return self.tools.zip_create(parts[0].strip(), parts[1].strip())
    def cmd_unzip(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: unzip <zip> :: <out_dir>"
        return self.tools.unzip(parts[0].strip(), parts[1].strip())

    def cmd_makepdf(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: makepdf <out.pdf> :: <text>"
        return self.tools.text_to_pdf(parts[0].strip(), parts[1].strip())
    def cmd_docx2pdf(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: docx2pdf <in.docx> :: <out.pdf>"
        return self.tools.docx_to_pdf(parts[0].strip(), parts[1].strip())
    def cmd_pdf2docx(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: pdf2docx <in.pdf> :: <out.docx>"
        return self.tools.pdf_to_docx(parts[0].strip(), parts[1].strip())
    def cmd_pdfmerge(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: pdfmerge <out.pdf> :: <a.pdf>;<b.pdf>"
        outp=parts[0].strip(); pdfs=[p.strip() for p in parts[1].split(";") if p.strip()]
        if not pdfs: return "Need input PDFs."
        return self.tools.pdf_merge(outp, *pdfs)
    def cmd_pdfsplit(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: pdfsplit <in.pdf> :: <out_dir>"
        return self.tools.pdf_split(parts[0].strip(), parts[1].strip())
    def cmd_pdfextract(self, arg:str)->str:
        parts=arg.split("::",2)
        if len(parts)!=3: return "Usage: pdfextract <in.pdf> :: <out.pdf> :: <start>-<end>"
        try:
            rng=parts[2].strip(); m=re.match(r"(\d+)\s*-\s*(\d+)", rng)
            if not m: return "Range?"
            s=int(m.group(1)); e=int(m.group(2))
            return self.tools.pdf_extract(parts[0].strip(), parts[1].strip(), s, e)
        except Exception: return "Range?"

    def cmd_docload(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: docload <session_id> :: <path>"
        return self.tools.doc_load(parts[0].strip(), parts[1].strip())
    def cmd_docs(self, arg:str)->str:
        items = self.mem.doc_list()
        if not items: return "No sessions."
        return "\n".join([f"{sid}: {meta.get('path','?')}" for sid,meta in items])
    def cmd_docask(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: docask <session_id> :: <question>"
        return self.tools.doc_ask(parts[0].strip(), parts[1].strip(), self.llm_concise)

    def cmd_clipget(self, arg:str)->str: return self.tools.clip_get()
    def cmd_clipset(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: clipset :: <text>"
        return self.tools.clip_set(parts[1].strip())
    def cmd_screenshot(self, arg:str)->str:
        if not arg.strip(): return "Usage: screenshot <out.png>"
        return self.tools.screenshot(arg.strip())

    def cmd_open(self, arg:str)->str:
        target = arg.strip().lower()
        KNOWN = [
            ("YouTube",           "https://www.youtube.com",      ["youtube", "yt", "you tube"]),
            ("Google Calendar",   "https://calendar.google.com",  ["google calendar", "gcal", "calendar"]),
            ("Gmail",             "https://mail.google.com",      ["gmail", "g mail"]),
            ("Google",            "https://www.google.com",       ["google"]),
            ("Google Drive",      "https://drive.google.com",     ["gdrive", "google drive", "drive"]),
            ("GitHub",            "https://github.com",           ["github", "git hub"]),
        ]
        for label, url, keys in KNOWN:
            for k in keys:
                if target == k or target.startswith(k):
                    webbrowser.open(url); return f"Opening {label}."
        if "." in target and " " not in target:
            if not target.startswith("http"): target = "https://" + target
            try:
                webbrowser.open(target); return "Opening website."
            except Exception:
                return "Open failed."
        try:
            subprocess.Popen(arg.strip(), shell=True); return "Opening application."
        except Exception:
            return "Open failed."

    def cmd_note(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: note [#tag #tag2] :: <text>"
        tags=[t[1:] for t in parts[0].strip().split() if t.startswith("#")]
        nid = self.notes.add(parts[1].strip(), tags); return f"Note {nid} added."
    def cmd_notes(self, arg:str)->str:
        q=arg.strip()
        ns=self.notes.list() if not q else self.notes.search(q)
        if not ns: return "No notes."
        lines=[]
        for n in ns[:30]:
            tags=(" ".join(f"#{t}" for t in n.get("tags",[]))).strip()
            lines.append(f"{n['id']} {tags} :: {n['text'][:90]}")
        return "\n".join(lines)
    def cmd_notesum(self, arg:str)->str: return self.notes.summary()

    def cmd_remind(self, arg:str)->str:
        parts=arg.split("::",1)
        if len(parts)!=2: return "Usage: remind <YYYY-MM-DD HH:MM> :: <text>"
        when=parts[0].strip(); text=parts[1].strip()
        try:
            due=datetime.strptime(when, "%Y-%m-%d %H:%M")
            rid=self.mem.add_reminder(text, due, None)
            return f"Reminder {rid} set."
        except Exception: return "Time format?"
    def cmd_reminders(self, arg:str)->str:
        rs=self.mem.list_reminders()
        if not rs: return "None."
        lines=[f"{r['id']} [{'on' if r.get('enabled') else 'off'}] {r['due']} :: {r['text']}" for r in rs]
        return "\n".join(lines)
    def cmd_remindon(self, arg:str)->str:
        rid=arg.strip(); 
        if not rid: return "Usage: remindon <id>"
        ok=self.mem.set_reminder_enabled(rid, True); return "Enabled." if ok else "Not found."
    def cmd_remindoff(self, arg:str)->str:
        rid=arg.strip(); 
        if not rid: return "Usage: remindoff <id>"
        ok=self.mem.set_reminder_enabled(rid, False); return "Disabled." if ok else "Not found."

    # Audio / TTS
    def cmd_audio(self, arg:str) -> str:
        a = (arg or "").strip().lower()
        if not a or a == "list":
            try:
                devs = sd.query_devices()
                lines = []
                defaults = sd.default.device
                for i, d in enumerate(devs):
                    mark = ""
                    if isinstance(defaults, (list, tuple)):
                        if i == defaults[0]: mark += " [default-in]"
                        if i == defaults[1]: mark += " [default-out]"
                    lines.append(f"{i:>2}: {d['name']}  in:{int(d['max_input_channels'])} out:{int(d['max_output_channels'])}{mark}")
                return "\n".join(lines) or "No devices."
            except Exception as e:
                return f"Audio list error: {e}"
        if a.startswith("in "):
            try:
                idx = int(a.split()[1])
                CONFIG["input_device"] = idx
                self._save_config()
                self.listener._prepare_device()
                self.listener.calibrate(0.6)
                return f"Input device set to {idx}."
            except Exception:
                return "Usage: audio in <index>"
        if a == "autotune":
            self.listener._prepare_device()
            self.listener.calibrate(0.6)
            return "Mic auto-tuned."
        return "Usage: audio [list|in <index>|autotune]"

    def cmd_rate(self, arg:str) -> str:
        try:
            rate = int(arg.strip())
            CONFIG["voice_rate"] = rate; self._save_config()
            return "Rate set."
        except Exception:
            return "Usage: rate <int>   (SAPI range -10..+10)"
    def cmd_volume(self, arg:str) -> str:
        try:
            vol = int(arg.strip())
            CONFIG["voice_volume"] = vol; self._save_config()
            return "Volume set."
        except Exception:
            return "Usage: volume <0..100>"
    def cmd_say(self, arg:str) -> str:
        msg = arg.strip() or "Testing one two three."
        self.speaker.say(msg)
        return "Speaking."

    # Math
    def cmd_calc(self, arg: str) -> str:
        if not arg.strip(): return "Usage: calc <expr>  |  example:  5 km to m"
        return self.calc.calc(arg.strip())
    def cmd_solve(self, arg: str) -> str:
        if not arg.strip(): return "Usage: solve <expr>  |  example:  solve x^2 - 5*x + 6"
        return self.math.solve_expr(arg.strip())
    def cmd_factor(self, arg: str) -> str:
        if not arg.strip(): return "Usage: factor <n>"
        return self.math.factor(arg.strip())
    def cmd_modinv(self, arg: str) -> str:
        parts = arg.split("::", 1)
        if len(parts) != 2: return "Usage: modinv <a> :: <m>"
        return self.math.mod_inv(parts[0].strip(), parts[1].strip())
    def cmd_crt(self, arg: str) -> str:
        parts = arg.split("::", 1)
        if len(parts) != 2: return "Usage: crt <m1,m2,...> :: <r1,r2,...>"
        try:
            mods = [int(x) for x in parts[0].strip().split(",")]
            rems = [int(x) for x in parts[1].strip().split(",")]
            if len(mods) != len(rems): return "Mismatch."
            return self.math.crt(mods, rems)
        except Exception: return "Invalid CRT input."
    def cmd_rsa(self, arg: str) -> str:
        parts = arg.split("::", 1)
        if len(parts) != 2: return "Usage: rsa <n> :: <e>"
        return self.math.rsa_from_n_e(parts[0].strip(), parts[1].strip())

    # Languages
    def cmd_lang(self, arg:str) -> str:
        a = (arg or "").strip().lower()
        if not a or a == "show":
            stt = CONFIG.get("whisper_language","en")
            rep = CONFIG.get("reply_language","en")
            return f"STT: {stt}   Reply: {rep}"
        if a.startswith("stt "):
            code = a.split(" ",1)[1].strip()
            CONFIG["whisper_language"] = code
            self._save_config()
            return f"STT language set to {code}."
        if a.startswith("reply "):
            code = a.split(" ",1)[1].strip()
            CONFIG["reply_language"] = code
            self.llm_concise.reply_lang = code
            self.llm_funny.reply_lang = code
            self._save_config()
            return f"Reply language set to {code}."
        return "Usage: lang show | lang stt <code|auto> | lang reply <code>"

    # Profile
    def cmd_callme(self, arg:str) -> str:
        name = arg.strip()
        if not name:
            return "Usage: callme <name>"
        self.mem.set_display_name(name)
        self.user_name = name
        return f"Okay, I'll call you {name}."
    def cmd_whoami(self, arg:str) -> str:
        if self.user_name:
            return f"You are {self.user_name}."
        return "I don't have a saved name for you yet. Say: call me <name>."

    # STT controls
    def cmd_stt(self, arg: str) -> str:
        a = (arg or "").strip().lower()
        if a.startswith("quality "):
            q = a.split(" ",1)[1].strip()
            mapping = {
                "best": "large-v3",
                "high": "medium.en",
                "balanced": "small.en",
                "fast": "base.en",
            }
            if q not in mapping:
                return "Use: stt quality <best|high|balanced|fast>"
            model = mapping[q]
            with Spinner(f"Loading STT model ({model})…"):
                msg = self.listener.set_model(model)
            self._save_config()
            return msg
        if a.startswith("model "):
            model = a.split(" ",1)[1].strip()
            with Spinner(f"Loading STT model ({model})…"):
                msg = self.listener.set_model(model)
            self._save_config()
            return msg
        if a == "reload":
            model = CONFIG.get("whisper_model","medium.en")
            with Spinner(f"Reloading STT model ({model})…"):
                msg = self.listener.set_model(model)
            return msg
        return "Usage: stt quality <best|high|balanced|fast> | stt model <name> | stt reload"

    # Debug
    def cmd_debug(self, arg: str) -> str:
        parts = (arg or "").strip().split()
        if len(parts) == 2 and parts[0] == "stt":
            try:
                secs = float(parts[1]); secs = max(0.5, min(10.0, secs))
            except Exception:
                return "Usage: debug stt <seconds>"
            rate = self.listener.rate
            n = int(rate * secs)
            try:
                buf = sd.rec(n, samplerate=rate, channels=1, dtype="float32", device=self.listener.device)
                sd.wait()
            except Exception as e:
                return f"Record error: {e}"
            raw_path = "debug_input.wav"
            with wave.open(raw_path, "wb") as wf:
                wf.setnchannels(1); wf.setsampwidth(2); wf.setframerate(rate)
                wf.writeframes((buf.squeeze()*32767).astype(np.int16).tobytes())
            arr16 = resample_to_16k(buf.squeeze(), rate)
            path16 = "debug_input_16k.wav"
            with wave.open(path16, "wb") as wf:
                wf.setnchannels(1); wf.setsampwidth(2); wf.setframerate(16000)
                wf.writeframes((arr16*32767).astype(np.int16).tobytes())
            try:
                lang = CONFIG.get("whisper_language","en")
                lang = None if (not lang or lang == "auto") else lang
                with self.listener.model_lock:
                    segments, _ = self.listener.model.transcribe(
                        arr16, language=lang,
                        beam_size=CONFIG.get("whisper_beam_size",5),
                        temperature=CONFIG.get("whisper_temperature",0.0),
                        vad_filter=False, without_timestamps=True, condition_on_previous_text=False
                    )
                text = " ".join(seg.text for seg in segments).strip()
            except Exception as e:
                text = f"[error: {e}]"
            return f"Saved {raw_path}, {path16}\nTranscript: {text or '[empty]'}"
        return "Usage: debug stt <seconds>"

    # ------------------ Intent helpers ------------------
    def _looks_time(self, s:str)->bool:
        t = s.lower()
        return any(kw in t for kw in [
            "what is the time", "what time is it", "current time", "time right now", "tell me the time"
        ])
    def _looks_factual(self, s:str)->bool:
        if not CONFIG.get("serpapi_key"): return False
        t = s.lower().strip()
        must_web = any(t.startswith(p) for p in ["who is", "who was", "what is the capital of", "president of", "who's"])
        if must_web: return True
        kws = ["who","what","when","where","why","how many","how much","current","today","capital","president","population","price","date","weather","news"]
        chit = ["who are you", "what can you do"]
        if any(x in t for x in chit): return False
        return any(t.startswith(k) or f" {k} " in t for k in kws)
    def _looks_chitchat(self, s:str)->bool:
        t = s.lower().strip()
        chit = ["hi","hello","hey","how are you","what's up","whats up","thank you","thanks","good morning","good night","tell me a joke","introduce yourself","nice to meet you","yes please","ok","okay","sure","can you hear me"]
        return any(x == t or x in t for x in chit) or (len(t) <= 60 and not self.math.is_mathy(t) and not self._looks_factual(t))
    def _looks_complicated(self, s:str)->bool:
        t = s.lower()
        long = len(t) > 120 or t.count(".") + t.count("?") + t.count("!") >= 2
        hard_kw = ["explain","analyze","design","write code","pseudocode","algorithm","optimize","prove","derivation","step by step","troubleshoot","debug","walkthrough"]
        return long or any(k in t for k in hard_kw)
    def _looks_callme(self, s:str) -> Optional[str]:
        t = s.strip()
        m = re.match(r"(?i)call\s+me(?:\s+as)?\s+(.+)$", t)
        if m:
            name = m.group(1).strip().strip(".")
            return name if name else None
        return None

    def _correct_open_target(self, raw: str) -> Optional[str]:
        known = {
            "google": ["google", "gogle", "gooogle"],
            "gmail": ["gmail", "g mail", "gee mail"],
            "youtube": ["youtube", "you tube", "yt", "utub", "u tube", "you tub"],
            "google calendar": ["google calendar", "gcal", "calendar", "calender", "g calendar"],
            "gdrive": ["gdrive", "google drive", "drive"],
            "github": ["github", "git hub"],
        }
        s = re.sub(r"[^a-z0-9]+", " ", raw.lower()).strip()
        for canon, aliases in known.items():
            if s in aliases or canon == s:
                return canon
        all_aliases = []
        for canon, aliases in known.items():
            for a in aliases:
                all_aliases.append((a, canon))
        candidates = difflib.get_close_matches(s, [a for a,_ in all_aliases], n=1, cutoff=0.72)
        if candidates:
            chosen = candidates[0]
            for a, canon in all_aliases:
                if a == chosen:
                    return canon
        return None

    # ------------------ Core run ------------------
    def _save_config(self):
        try:
            with open(CONFIG["config_file"],"w",encoding="utf-8") as f:
                json.dump(CONFIG,f,indent=2,ensure_ascii=False)
        except Exception as e:
            logger.warning("Config save failed: %s", e)
    def _load_config(self):
        try:
            if os.path.exists(CONFIG["config_file"]):
                with open(CONFIG["config_file"],"r",encoding="utf-8") as f:
                    conf=json.load(f); CONFIG.update(conf)
        except Exception as e:
            logger.warning("Config load failed: %s", e)

    def start(self):
        if CONFIG.get("console_banner", True):
            os.system("cls" if os.name=="nt" else "clear")
            print("🟢 Jarvis Prime v6 — speak anytime; type commands.  (say 'exit' to quit)")
        self.speaker.say("Jarvis online. Ready.")
        self.listener.start()
        self.typist.start()
        self.scheduler.start()
        self.running=True
        try:
            while self.running:
                if self.listener.thread and not self.listener.thread.is_alive() and self.listener.running:
                    print("⚠️ Mic thread died — restarting…")
                    self.listener.stop(); time.sleep(0.2); self.listener.start()
                try:
                    src, text = self._next_input()
                    self._handle(src, text)
                except Exception as e:
                    logger.exception("Main loop error: %s", e); time.sleep(0.2)
        except KeyboardInterrupt:
            pass
        finally:
            print("\n👋 Shutting down.")
            self.speaker.say("Shutting down.")
            time.sleep(0.45)
            self.listener.stop(); self.typist.stop(); self.scheduler.stop(); self.speaker.stop()

    def _next_input(self) -> Tuple[str,str]:
        if not self.typist.out_q.empty(): return self.typist.out_q.get_nowait()
        if not self.listener.out_q.empty(): return self.listener.out_q.get_nowait()
        time.sleep(0.05); return ("", "")

    def _handle(self, src:str, text:str):
        t = text.strip()
        if not t: return

        # Heuristic correction: one‑word opens ("google", "gmail", ...)
        if t.lower() in {"google","gmail","youtube","calendar","drive","github"}:
            t = f"open {t.lower()}"

        # Auto-correct "open ..." targets
        if t.lower().startswith("open "):
            rest = t[5:].strip()
            fixed = self._correct_open_target(rest)
            if fixed:
                t = f"open {fixed}"

        print(f"\n🎧 ({src or 'sys'}) {t}")
        logger.info("IN (%s): %s", src, t)

        wanted_name = self._looks_callme(t)
        if wanted_name:
            self.mem.set_display_name(wanted_name)
            self.user_name = wanted_name
            self._respond(f"Okay, I'll call you {wanted_name}.", "sys", t)
            return

        res = self.commands.resolve(t)
        if res:
            name, meta = res
            arg = t[len(name):].strip() if t.lower().startswith(name) else (t.split(" ",1)[1] if " " in t else "")
            try: reply = meta["func"](arg)
            except Exception as e:
                logger.error("Command %s error: %s", name, e); reply = "Error."
            self._respond(reply, name, t); return

        if self._looks_time(t):
            now = datetime.now().strftime("%I:%M %p").lstrip("0")
            self._respond(f"It's {now}.", "sys", t); return

        # Solver → calc → LLM for math
        if self.math.is_mathy(t):
            reply = self.math.solve_expr(t)
            def _bad(x: str) -> bool: return x.strip().lower() in ("invalid", "invalid.")
            if _bad(reply):
                alt = self.calc.calc(t)
                if _bad(alt):
                    with Spinner("Thinking…"):
                        reply = self.llm_concise.ask(t, history=list(self.history), reply_lang=CONFIG.get("reply_language","en"), user_name=self.user_name)
                    self._respond(reply, "llm", t); return
                else:
                    self._respond(alt, "solver", t); return
            else:
                self._respond(reply, "solver", t); return

        # Web factual first
        if self._looks_factual(t):
            with Spinner("Thinking…"):
                resw = self.tools.web_lookup(t)
            if resw["confidence"] >= 0.55:
                self._respond(resw["text"], "web", t); return
            else:
                with Spinner("Thinking…"):
                    reply = self.llm_concise.ask(t, history=list(self.history), reply_lang=CONFIG.get("reply_language","en"), user_name=self.user_name)
                self._respond(reply, "llm", t); return

        # Chit-chat / short → funny persona
        if self._looks_chitchat(t) and not self._looks_complicated(t):
            with Spinner("Thinking…"):
                reply = self.llm_funny.ask(t, history=list(self.history), reply_lang=CONFIG.get("reply_language","en"), user_name=self.user_name)
            self._respond(reply, "llm", t); return

        # Otherwise concise LLM
        with Spinner("Thinking…"):
            reply = self.llm_concise.ask(t, history=list(self.history), reply_lang=CONFIG.get("reply_language","en"), user_name=self.user_name)
        self._respond(reply, "llm", t)

    def _respond(self, reply:str, source:str, user_text:str):
        if len(reply) > CONFIG["max_chars"]:
            reply = reply[:CONFIG["max_chars"]].rsplit(" ",1)[0] + "..."
        print(f"💬 [{source}] {reply}")
        logger.info("OUT (%s): %s", source, reply)
        self.history.append({"role":"user", "text": user_text})
        self.history.append({"role":"assistant", "text": reply})
        self.speaker.say(reply)
        self.mem.log(user_text, reply, source)

# ------------------------------------------------------------
if __name__ == "__main__":
    JarvisPrime().start()
